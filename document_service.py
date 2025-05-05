"""
This module handles document generation and export functionality.
"""

import os
from datetime import datetime
import docx
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

class DocumentService:
    """
    Provides document generation and export functionality
    """
    
    @staticmethod
    def save_text(content, file_path):
        """
        Save text content to a file
        
        Args:
            content (str): Text content to save
            file_path (str): Path to save the file
            
        Returns:
            tuple: (success, error_message)
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True, ""
        except Exception as e:
            return False, str(e)
    
    @staticmethod
    def export_word(transcript, analysis, output_path):
        """
        Export transcript and analysis to a Word document
        
        Args:
            transcript (str): The transcript text
            analysis (str): The analysis text
            output_path (str): Path to save the Word document
            
        Returns:
            tuple: (success, error_message)
        """
        try:
            # Create a new document
            doc = docx.Document()
            
            # Configure document styles
            style_normal = doc.styles['Normal']
            style_normal.font.size = Pt(11)
            
            # Add title
            doc.add_heading('Transcript and Analysis', 0)
            
            # Add timestamp
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add transcript section
            doc.add_heading('Original Transcript', 1)
            doc.add_paragraph(transcript)
            
            # Add analysis section
            doc.add_heading('Analysis', 1)
            
            # Process full analysis content with Markdown-like formatting
            analysis_lines = analysis.split('\n')
            
            current_paragraph = None
            
            for line in analysis_lines:
                # Handle level 2 headers (##)
                if line.startswith('## '):
                    doc.add_heading(line.replace('## ', ''), 2)
                
                # Handle level 3 headers (###)
                elif line.startswith('### '):
                    doc.add_heading(line.replace('### ', ''), 3)
                
                # Handle bullet points
                elif line.strip().startswith('- '):
                    bullet_text = line.strip().replace('- ', '', 1)
                    doc.add_paragraph(bullet_text, style='List Bullet')
                
                # Regular paragraphs or empty lines
                elif not line.strip():
                    current_paragraph = None  # Start a new paragraph after empty line
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph(line)
                    else:
                        current_paragraph.add_run('\n' + line)
            
            # Save the document
            doc.save(output_path)
            return True, ""
            
        except Exception as e:
            return False, str(e)
    
    @staticmethod
    def export_pdf(transcript, analysis, output_path):
        """
        Export transcript and analysis to a PDF document
        
        Args:
            transcript (str): The transcript text
            analysis (str): The analysis text
            output_path (str): Path to save the PDF document
            
        Returns:
            tuple: (success, error_message)
        """
        try:
            # Create a PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = styles['Title']
            heading1_style = styles['Heading1']
            heading2_style = styles['Heading2']
            heading3_style = styles['Heading3']
            normal_style = styles['Normal']
            normal_style.fontSize = 11
            
            # Create bullet style
            bullet_style = ParagraphStyle(
                'BulletPoint',
                parent=normal_style,
                leftIndent=20,
                firstLineIndent=0,
                spaceBefore=2,
                bulletIndent=10,
                bulletFontName='Symbol',
                bulletFontSize=11
            )
            
            # Create content elements
            elements = []
            
            # Add title
            elements.append(Paragraph('Transcript and Analysis', title_style))
            elements.append(Spacer(1, 12))
            
            # Add timestamp
            elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
            elements.append(Spacer(1, 20))
            
            # Add transcript section
            elements.append(Paragraph('Original Transcript', heading1_style))
            elements.append(Spacer(1, 12))
            formatted_transcript = transcript.replace('\n', '<br />')
            elements.append(Paragraph(formatted_transcript, normal_style))
            elements.append(Spacer(1, 20))
            
            # Add analysis section
            elements.append(Paragraph('Analysis', heading1_style))
            elements.append(Spacer(1, 12))
            
            # Process full analysis content with Markdown-like formatting
            analysis_lines = analysis.split('\n')
            current_text = ""
            
            for line in analysis_lines:
                # Handle level 2 headers (##)
                if line.startswith('## '):
                    if current_text:
                        elements.append(Paragraph(current_text, normal_style))
                        current_text = ""
                    elements.append(Spacer(1, 8))
                    elements.append(Paragraph(line.replace('## ', ''), heading2_style))
                    elements.append(Spacer(1, 4))
                
                # Handle level 3 headers (###)
                elif line.startswith('### '):
                    if current_text:
                        elements.append(Paragraph(current_text, normal_style))
                        current_text = ""
                    elements.append(Spacer(1, 6))
                    elements.append(Paragraph(line.replace('### ', ''), heading3_style))
                    elements.append(Spacer(1, 3))
                
                # Handle bullet points
                elif line.strip().startswith('- '):
                    if current_text:
                        elements.append(Paragraph(current_text, normal_style))
                        current_text = ""
                    bullet_text = line.strip().replace('- ', '', 1)
                    elements.append(Paragraph(f"• {bullet_text}", bullet_style))
                
                # Regular paragraphs
                elif line.strip():
                    if current_text:
                        current_text += "<br />"
                    current_text += line
                
                # Empty lines
                elif current_text:
                    elements.append(Paragraph(current_text, normal_style))
                    current_text = ""
                    elements.append(Spacer(1, 6))
            
            # Add any remaining text
            if current_text:
                elements.append(Paragraph(current_text, normal_style))
            
            # Build the PDF document
            doc.build(elements)
            return True, ""
            
        except Exception as e:
            return False, str(e) 